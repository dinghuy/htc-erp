from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_PATH = r"C:\Users\dinghuy\OneDrive - HUYNH THY GROUP\Antigravity Workspace\htc-erp\docs\bao-cao-tom-tat-he-thong-minierp-2026-03-31-v2.docx"
PRIMARY = "1F4E78"


def set_font(run, size=10.5, bold=False, color=None, italic=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, *, bold=False, size=10.0, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table, widths):
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Cm(widths[idx])
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(3)


def add_title(document, title_text, subtitle_text):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title_text)
    set_font(r, size=16, bold=True)
    s = document.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run(subtitle_text)
    set_font(sr, size=10.5, italic=True)


def add_heading(document, text):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, size=13, bold=True, color=PRIMARY)


def add_paragraph(document, text):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r)


def add_bullets(document, items):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_font(r)


def add_table(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    style_table(table, widths)
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color="FFFFFF")
        shade_cell(table.rows[0].cells[idx], PRIMARY)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            set_cell_text(cells[idx], value)
    return table


def build_document():
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)

    add_title(
        document,
        "BÁO CÁO TÓM TẮT HỆ THỐNG miniERP / CRM HUỲNH THY",
        "Update đến ngày 31/03/2026",
    )

    add_heading(document, "1. Hệ thống này hiện làm được gì")
    add_bullets(
        document,
        [
            "Quản lý thông tin khách hàng, người liên hệ và đầu mối bán hàng trên cùng một hệ thống.",
            "Lưu danh mục sản phẩm, nhà cung cấp và giá đầu vào để làm báo giá nhanh hơn và giảm lệ thuộc vào file Excel rời rạc.",
            "Gom lại, lưu trữ các tài liệu, hình ảnh, video liên quan đến sản phẩm để xem nhanh và chia sẻ thông tin cho khách hàng.",
            "Tạo báo giá, lưu lịch sử báo giá và xuất file PDF để gửi khách hàng.",
            "Theo dõi quá trình từ báo giá sang đơn hàng, mua hàng, hàng về và giao/bàn giao ở mức có thể trình diễn và dùng thử nội bộ.",
            "Gom các việc cần làm, việc chờ duyệt và các vấn đề còn thiếu vào cùng hệ thống để các bộ phận dễ phối hợp hơn.",
            "Quản lý người dùng, phân quyền cơ bản và lưu lại lịch sử thao tác để dễ kiểm soát.",
        ],
    )

    add_heading(document, "2. Hệ thống này giúp ích gì cho công ty")
    add_bullets(
        document,
        [
            "Rút ngắn thời gian làm báo giá và giảm việc copy thủ công giữa nhiều file.",
            "Giúp các bộ phận nhìn chung một dòng công việc, thay vì mỗi bộ phận giữ một file riêng.",
            "Làm rõ tình trạng một đơn hàng đang ở bước nào: đang báo giá, đang mua hàng, đang chờ hàng về hay đang chuẩn bị giao.",
            "Giúp theo dõi những việc đang chờ duyệt để tránh tình trạng chậm xử lý hoặc tồn đọng ở một bước nào đó.",
            "Giúp quản lý theo dõi được tình hình công việc và giảm rủi ro bỏ sót thông tin.",
            "Tạo nền để sau này nhìn được nhiều dự án cùng lúc, biết ai đang xử lý việc gì và khối lượng công việc đang dồn ở đâu.",
            "Tạo nền để sau này nối liền công việc từ Sales sang Mua hàng, Kho, Kế toán và Vận hành theo một quy trình chung.",
        ],
    )

    add_heading(document, "3. Những tính năng đã hoàn thiện hoặc đã có nền")
    add_table(
        document,
        ["Nhóm chức năng", "Mức hiện tại", "Ý nghĩa thực tế"],
        [
            ("Quản lý khách hàng / liên hệ / đầu mối", "Đã làm xong để dùng thử", "Giúp Sales quản lý đầu vào bán hàng trên hệ thống."),
            ("Quản lý sản phẩm / thiết bị", "Đã làm xong để dùng thử", "Là dữ liệu nền để làm báo giá đúng và nhanh hơn."),
            ("Quản lý nhà cung cấp / giá đầu vào", "Đã làm xong để dùng thử", "Giúp kiểm soát giá vốn và chuẩn bị cho bước mua hàng."),
            ("Tài liệu / hình ảnh / video sản phẩm", "Đã làm xong để dùng thử", "Giúp lưu thông tin sản phẩm tập trung để xem nhanh và gửi cho khách hàng."),
            ("Báo giá / xuất PDF / theo dõi trạng thái", "Đã làm xong để dùng thử", "Là phần nổi bật nhất hiện nay, có thể dùng để trình bày với quản lý và khách nội bộ."),
            ("Việc cần làm / việc chờ duyệt / việc cần bổ sung", "Đã có nền để demo và hoàn thiện tiếp", "Giúp gom những việc đang chờ xử lý vào một chỗ để dễ theo dõi."),
            ("Người dùng / phân quyền / cài đặt cơ bản", "Đã làm xong để dùng thử", "Giúp kiểm soát ai được xem và ai được thao tác."),
            ("Báo cáo cơ bản / lịch sử thao tác", "Đã có nền chạy được", "Hỗ trợ quản lý xem tình hình chung và theo dõi thao tác."),
            ("Đơn hàng / mua hàng / tiến độ hàng về / giao hàng", "Đã có nền để demo và hoàn thiện tiếp", "Là phần đang được nối tiếp từ báo giá sang vận hành thực tế."),
            ("Task / Project / điều phối công việc", "Đã có nền để mở rộng", "Dùng để theo dõi các bước sau khi đơn hàng được chốt."),
            ("Quản lý nhiều dự án / nhìn tiến độ chung", "Đã có nền để mở rộng", "Tạo nền để quản lý nhìn tổng quan nhiều dự án và phân bổ công việc hợp lý hơn."),
        ],
        [5.0, 4.0, 8.0],
    )

    add_heading(document, "4. Hiện hệ thống đang vận hành theo cách nào")
    add_bullets(
        document,
        [
            "Mỗi người dùng sẽ thấy các màn hình phù hợp với vai trò của mình, ví dụ Sales, Mua hàng, Kế toán hoặc quản lý.",
            "Các bộ phận cùng nhìn vào một công việc chung, thay vì tách rời theo từng file và từng người giữ thông tin riêng.",
            "Hệ thống đã có phần gom việc theo kiểu việc cần làm, việc chờ duyệt, việc bị thiếu hồ sơ hoặc đang bị vướng ở một bước nào đó.",
            "Dữ liệu được lưu tập trung trong hệ thống nội bộ, phù hợp với định hướng kiểm soát dữ liệu của công ty.",
            "Cách phát triển hiện nay cũng đã rõ hơn: làm theo từng hạng mục nhỏ, kiểm tra lại, rồi mới đưa vào dùng thử.",
        ],
    )

    add_heading(document, "5. Hướng đi tiếp theo")
    add_table(
        document,
        ["Việc tiếp theo", "Mục tiêu"],
        [
            ("Tiếp tục hoàn thiện đoạn từ báo giá đến giao hàng", "Để hệ thống không chỉ làm báo giá mà còn theo dõi được tiếp các bước sau bán hàng."),
            ("Cho các bộ phận dùng thử nội bộ", "Để kiểm tra đúng/sai, xem bước nào còn vướng và điều chỉnh theo thực tế làm việc."),
            ("Chuẩn hóa quy trình chung giữa Sales, Mua hàng, Kế toán, Kho và Vận hành", "Để sau này phát triển thành hệ thống ERP tổng thể, không bị mỗi bộ phận một cách làm."),
            ("Tiếp tục hoàn thiện các phần đã có nền", "Để biến phần đang demo được thành phần có thể dùng ổn định hơn trong thực tế."),
            ("Bổ sung phần duyệt và đo thời gian chờ duyệt", "Để thấy bước nào đang bị nghẽn, xử lý nhanh hơn và có dữ liệu để đánh giá hiệu quả phối hợp."),
            ("Quản lý được nhiều dự án cùng lúc và nhìn khối lượng công việc theo người", "Để quản lý biết nhân sự nào đang quá tải, dự án nào đang chậm và điều phối nguồn lực tốt hơn."),
        ],
        [6.2, 10.8],
    )

    add_heading(document, "6. Kết luận ngắn")
    add_paragraph(
        document,
        "Tính đến hiện tại, hệ thống không còn là ý tưởng trên giấy mà đã có nhiều phần dùng được và đủ để cho thấy hướng phát triển rõ ràng. Phần mạnh nhất hiện nay là quản lý dữ liệu bán hàng, làm báo giá và tạo nền để nối tiếp sang đơn hàng, mua hàng và giao hàng. Giai đoạn tiếp theo là làm cho đoạn vận hành sau báo giá chạy mượt hơn để tiến tới một hệ thống dùng chung cho nhiều bộ phận.",
    )

    return document


if __name__ == "__main__":
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)
