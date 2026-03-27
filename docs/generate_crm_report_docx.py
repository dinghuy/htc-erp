from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_PATH = r"C:\Users\dinghuy\OneDrive - HUYNH THY GROUP\Antigravity Workspace\crm-app\docs\bao-cao-ke-hoach-crm-huynh-thy-demo-go-live-v5.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 10.5, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table, widths_cm):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[idx])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.space_before = Pt(0)


def add_heading(document: Document, text: str, level: int = 1):
    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(4)
    run = heading.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(14 if level == 1 else 12)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)


def add_two_col_table(document: Document, rows, widths=(5.2, 11.8)):
    table = document.add_table(rows=1, cols=2)
    style_table(table, widths)
    header = table.rows[0].cells
    set_cell_text(header[0], "Hạng mục", bold=True, color="FFFFFF")
    set_cell_text(header[1], "Nội dung", bold=True, color="FFFFFF")
    set_cell_shading(header[0], "1F4E78")
    set_cell_shading(header[1], "1F4E78")
    for left, right in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], left, bold=True)
        set_cell_text(cells[1], right)
    return table


def add_multi_col_table(document: Document, headers, rows, widths_cm):
    table = document.add_table(rows=1, cols=len(headers))
    style_table(table, widths_cm)
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color="FFFFFF")
        set_cell_shading(table.rows[0].cells[idx], "1F4E78")
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            set_cell_text(cells[idx], value)
    return table


def money(value: str) -> str:
    return value


def build_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("BÁO CÁO KẾ HOẠCH TRIỂN KHAI CRM HUỲNH THY")
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Phiên bản báo cáo phục vụ Demo Sales-to-Delivery cuối tháng 4/2026 và Go-live nội bộ sau demo")
    sub_run.italic = True
    sub_run.font.name = "Arial"
    sub_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    sub_run.font.size = Pt(10.5)

    add_heading(document, "1. Mục tiêu demo / pilot")
    add_two_col_table(document, [
        ("Ngày lập báo cáo", "25/03/2026"),
        ("Mục tiêu mốc 1", "Hoàn thành bản demo dạng mini app chạy thử được theo luồng end-to-end Sales-to-Delivery vào ngày 30/04/2026."),
        ("Mục tiêu mốc 2", "Sau demo, tiếp tục chạy nội bộ để rà soát tính đúng/sai, hiệu quả sử dụng và cải tiến thêm."),
        ("Nhóm dùng đầu tiên", "Sales là nhóm dùng trước, nhưng bản demo phải bao phủ handoff sang Mua hàng, Kho và Delivery."),
        ("Trục triển khai", "Sales-to-Delivery: từ lead/báo giá đến đơn hàng, mua hàng, hàng về kho và giao/bàn giao."),
        ("Nguồn lực hiện tại", "Huy SA (lead tổng thể, logic, AI coding), Huy IT (database, hạ tầng, deploy)."),
        ("Quy mô công ty", "Khoảng 50 người; giai đoạn đầu đề xuất pilot 5-10 người dùng thực tế."),
    ])

    add_heading(document, "2. Hiện trạng tính năng đã triển khai")
    add_multi_col_table(
        document,
        ["Nhóm tính năng", "Hiện trạng", "Vai trò trong phase demo/pilot"],
        [
            ("Đăng nhập / đổi mật khẩu / phân quyền", "Đã triển khai", "Là nền bắt buộc để chạy thử nội bộ và kiểm soát user theo vai trò."),
            ("Accounts / Contacts / Leads", "Đã triển khai", "Là đầu vào chính cho luồng Sales-to-Quotation."),
            ("Products / Equipment", "Đã triển khai", "Cung cấp dữ liệu sản phẩm và thông số cho báo giá."),
            ("Suppliers / Supplier Quotes", "Đã triển khai", "Là nền để kiểm soát giá đầu vào và phục vụ bước mua hàng."),
            ("Pricing / tỷ giá / VAT / điều khoản", "Đã triển khai", "Là lõi tính toán để chốt giá và chuẩn hóa báo giá."),
            ("Quotation / PDF / status", "Đã triển khai", "Là lõi trình diễn của đoạn Sales-to-Quotation."),
            ("Sales Orders / ERP handoff", "Đã có nền", "Là nền để đẩy luồng từ báo giá sang đơn hàng nội bộ."),
            ("Projects / Tasks / workflow pack", "Đã có nền", "Có thể tận dụng để biểu diễn các bước mua hàng, chuẩn bị giao và bàn giao."),
            ("Ops Overview / Gantt / Staff", "Đã có nền", "Là nền để nhìn theo dõi tiến độ thực hiện, hàng về và giao hàng."),
            ("Reports / Activity log / Dashboard", "Đã triển khai", "Hỗ trợ quản lý nhìn được tiến độ, funnel, trạng thái và lịch sử thao tác."),
            ("Settings / company info / i18n / Users", "Đã triển khai", "Hỗ trợ vận hành nội bộ, chuẩn hóa thông tin công ty và tài khoản."),
            ("Chat / Notifications", "Đã có nền", "Chưa là trọng tâm của demo nhưng có thể dùng khi mở rộng vận hành."),
            ("SQLite persistence / import-export", "Đã triển khai", "Phù hợp định hướng chạy nội bộ, kiểm soát dữ liệu và hoạt động trong LAN."),
        ],
        [4.8, 2.7, 9.0],
    )

    add_heading(document, "3. Timeline chính")
    add_multi_col_table(
        document,
        ["Giai đoạn", "Thời gian", "Đầu ra chính", "Phụ trách chính"],
        [
            ("Chốt scope end-to-end", "25/03 - 31/03", "Khóa phạm vi demo Sales-to-Delivery, chốt dữ liệu mẫu, trạng thái nghiệp vụ và các điểm handoff.", "Huy SA"),
            ("Hoàn thiện Sales-to-Quotation", "01/04 - 10/04", "Hoàn thiện end-to-end: Khách hàng -> Sản phẩm -> Pricing -> Báo giá -> tạo đơn hàng.", "Huy SA"),
            ("Mở rộng Order-to-Procurement / Delivery", "11/04 - 17/04", "Bổ sung trạng thái mua hàng, hàng về, chuẩn bị giao, bàn giao/giao hàng; hoàn thiện dữ liệu và UI liên quan.", "Huy SA"),
            ("UAT liên phòng ban", "18/04 - 24/04", "Test thực tế với Sales, Kế toán, Mua hàng, Kho/Vận hành; chỉnh luồng end-to-end và chốt kịch bản demo.", "Huy SA + các phòng"),
            ("Feature freeze & demo prep", "25/04 - 29/04", "Khóa tính năng, test hồi quy, dữ liệu demo, phương án fallback.", "Huy SA"),
            ("Demo quản lý", "30/04", "Trình diễn mini app theo luồng Sales-to-Delivery: báo giá -> đơn hàng -> mua hàng -> hàng về -> giao/bàn giao.", "Huy SA"),
            ("Chạy thử nội bộ", "01/05 - 10/05", "Team IT triển khai môi trường nội bộ, chạy thử nhóm nhỏ theo luồng Sales-to-Delivery.", "Huy IT + Huy SA"),
            ("Ổn định sau pilot", "11/05 - 25/05", "Fix theo feedback, đo hiệu quả end-to-end, chuẩn hóa báo cáo và mở rộng vận hành thực tế.", "Huy SA + Huy IT"),
        ],
        [4.0, 3.2, 8.0, 2.8],
    )

    add_heading(document, "4. Cơ chế phối hợp liên phòng ban")
    add_multi_col_table(
        document,
        ["Thành phần", "Vai trò", "Trách nhiệm trong phase đầu"],
        [
            ("Sales", "Đầu vào nhu cầu khách hàng", "Cung cấp quy trình báo giá, trạng thái deal, điểm nghẽn khi theo dõi khách hàng và báo giá."),
            ("Kế toán", "Kiểm soát số liệu thương mại", "Chốt VAT, điều khoản thanh toán, quy tắc số liệu và điều kiện xác nhận doanh thu/đơn hàng."),
            ("Mua hàng", "Kiểm soát đầu vào supplier", "Cung cấp giá đầu vào, lead time, điều kiện mua và điểm nghẽn với nhà cung cấp."),
            ("Kho / Vận hành", "Kiểm soát hàng về và giao hàng", "Xác nhận trạng thái hàng, nhập/xuất kho, chuẩn bị giao, bàn giao và nghiệm thu."),
            ("Huy SA", "Điều phối nghiệp vụ và hệ thống", "Tổng hợp yêu cầu liên phòng ban, mô hình hóa thành quy trình chung, chuyển hóa thành backlog hệ thống."),
            ("Sếp", "Người chốt cuối cùng", "Ra quyết định khi có khác biệt hoặc xung đột quy trình giữa các phòng ban."),
        ],
        [3.0, 3.2, 11.8],
    )

    add_heading(document, "5. Phương thức làm việc liên phòng ban")
    add_multi_col_table(
        document,
        ["Nội dung", "Cách thực hiện", "Đầu ra bắt buộc"],
        [
            ("Thu thập yêu cầu", "Phỏng vấn ngắn từng phòng ban theo cùng một khung câu hỏi.", "Biết rõ đầu vào, đầu ra, phê duyệt, file dùng, lỗi lặp lại và điểm còn làm tay/Excel."),
            ("Xác nhận quy trình", "Sau phỏng vấn, chốt lại bằng bản xác nhận ngắn cho từng phòng.", "Mỗi phòng xác nhận bước quy trình, dữ liệu nhập/xuất, trách nhiệm và điều kiện bàn giao."),
            ("Chuẩn hóa quy trình", "So sánh khác biệt giữa các phòng và gom về một quy trình chung.", "Có 1 quy trình thống nhất để đưa vào hệ thống, không để mỗi phòng 1 cách làm riêng."),
            ("Điều hành hàng tuần", "Họp cố định hàng tuần với đại diện Sales, Kế toán, Mua hàng, Kho/Vận hành.", "Mỗi buổi phải ra quyết định, người chịu trách nhiệm và backlog hệ thống tuần tới."),
            ("Triển khai thực tế", "Rollout từng bước theo luồng Sales-to-Delivery, không mở đồng loạt toàn công ty.", "Giảm rủi ro, dễ kiểm tra đúng/sai và sửa nhanh trong pilot."),
        ],
        [3.4, 5.2, 10.0],
    )

    add_heading(document, "6. Lộ trình ERP tổng thể")
    add_multi_col_table(
        document,
        ["Giai đoạn", "Phạm vi", "Mục tiêu quản trị"],
        [
            ("Sales-to-Quotation", "Lead, account, contact, product, pricing, quotation.", "Chuẩn hóa đầu vào khách hàng và đầu ra báo giá; rút ngắn thời gian xử lý báo giá."),
            ("Quotation-to-Order", "Handoff sang đơn hàng, điều kiện thanh toán, giá vốn, kiểm soát chốt đơn.", "Biến báo giá thành đơn hàng có kiểm soát thay vì dừng ở tài liệu báo giá."),
            ("Order-to-Procurement / Delivery", "Mua hàng, PO, theo dõi hàng về, nhập/xuất kho, giao/bàn giao.", "Chốt được một lát cắt ERP end-to-end ngay trong demo tháng 4, sau đó mới tối ưu và mở rộng sâu hơn."),
        ],
        [3.5, 5.8, 9.3],
    )

    add_heading(document, "7. Chi phí AI và công cụ hỗ trợ")
    add_multi_col_table(
        document,
        ["Nhóm chi phí", "Cơ sở tính", "Ước tính", "Ghi chú"],
        [
            ("AI triển khai - mức hiện tại", "300.000 VND/ngày; 25/03 - 25/05 (~62 ngày)", money("18.600.000 VND"), "Mức kiểm soát chi phí hiện tại."),
            ("AI triển khai - tăng tốc nếu cần", "500.000 - 800.000 VND/ngày; giai đoạn nước rút", money("31.000.000 - 49.600.000 VND"), "Chỉ dùng khi cần rút ngắn backlog trước demo."),
            ("Hạ tầng nội bộ hiện có", "Tận dụng tài nguyên nội bộ do Team IT triển khai", money("0 VND trong báo cáo giai đoạn này"), "Chưa tách chi phí hạ tầng riêng ở phase demo/pilot."),
            ("Công cụ hỗ trợ khác", "Phát sinh nếu có", money("Theo thực tế"), "Chỉ mở khi cần hỗ trợ trực tiếp cho triển khai."),
        ],
        [4.4, 5.4, 3.7, 4.5],
    )

    add_heading(document, "8. Rủi ro và kiến nghị")
    add_multi_col_table(
        document,
        ["Nội dung", "Tác động / mục tiêu", "Cách xử lý / kiến nghị"],
        [
            ("Scope phát sinh thêm trước demo", "Làm trễ critical path", "Khóa phạm vi phase tháng 4 vào đúng lát cắt Sales-to-Delivery; mọi yêu cầu mới ngoài flow này chuyển phase sau."),
            ("Dữ liệu đầu vào chưa chuẩn", "Demo/pilot dùng không mượt", "Làm sạch dữ liệu mẫu sớm, chốt format import ngay từ đầu tháng 4."),
            ("Chậm triển khai hạ tầng nội bộ", "Không kịp chạy thử sau demo", "Team IT chuẩn bị sớm môi trường local/network nội bộ song song với giai đoạn hoàn thiện core flow."),
            ("Thiếu UAT từ Sales/Kế toán/Mua hàng/Kho", "Sai lệch nghiệp vụ thực tế", "Cử user đại diện test theo lịch cố định trong tuần 18/04 - 24/04."),
            ("Rủi ro dữ liệu/kết nối bên ngoài", "Ảnh hưởng an toàn dữ liệu và vận hành", "Không dùng cloud bên ngoài ở phase đầu; ưu tiên mạng nội bộ, local DB và khả năng hoạt động khi mất internet."),
            ("Kiến nghị quản trị", "Giữ tốc độ và giảm xung đột liên phòng ban", "Huy SA tổng hợp yêu cầu và đề xuất phương án; sếp chốt cuối cùng khi có khác biệt giữa các phòng."),
        ],
        [4.2, 3.2, 10.6],
    )

    document.add_page_break()

    add_heading(document, "Phụ lục A. Bảng trách nhiệm theo phòng ban")
    add_multi_col_table(
        document,
        ["Phòng ban", "Đầu vào cần cung cấp", "Đầu ra cần xác nhận"],
        [
            ("Sales", "Nhu cầu khách hàng, trạng thái deal, thông tin account/contact.", "Nội dung báo giá, tình trạng chào giá, phản hồi khách hàng."),
            ("Kế toán", "Quy tắc VAT, điều khoản thanh toán, nguyên tắc kiểm soát số liệu.", "Xác nhận điều kiện tài chính đủ chuẩn để chốt đơn."),
            ("Mua hàng", "Giá supplier, lead time, điều kiện mua.", "Xác nhận đầu vào giá vốn và khả năng cung ứng."),
            ("Kho / Vận hành", "Trạng thái hàng, điều kiện nhập/xuất, kế hoạch giao hàng.", "Xác nhận sẵn sàng giao, bàn giao và nghiệm thu."),
        ],
        [3.0, 7.0, 8.6],
    )

    add_heading(document, "Phụ lục B. Bảng phân vai liên phòng ban theo lộ trình ERP")
    add_multi_col_table(
        document,
        ["Giai đoạn", "Sales", "Kế toán", "Mua hàng", "Kho / Vận hành", "Huy SA / Sếp"],
        [
            (
                "Sales-to-Quotation",
                "Tiếp nhận nhu cầu khách hàng, cập nhật lead/account/contact, chốt nội dung báo giá.",
                "Xác nhận quy tắc VAT, điều khoản thanh toán, nguyên tắc số liệu thương mại.",
                "Cung cấp dữ liệu supplier và đầu vào giá nếu cần tham chiếu.",
                "Xác nhận điều kiện giao hàng, thời gian đáp ứng và yêu cầu vận hành nếu có.",
                "Huy SA chuẩn hóa form và logic; sếp chốt phạm vi phase đầu.",
            ),
            (
                "Quotation-to-Order",
                "Xác nhận báo giá thắng, tạo yêu cầu chuyển sang đơn hàng, bàn giao đủ thông tin khách và điều kiện bán.",
                "Kiểm tra điều kiện tài chính, công nợ, đặt cọc, điều kiện đủ để tạo đơn hợp lệ.",
                "Xác nhận giá vốn, nguồn cung, lead time, điều kiện mua tương ứng với đơn hàng.",
                "Xác nhận yêu cầu giao hàng, thời điểm cần hàng, điều kiện bàn giao nội bộ.",
                "Huy SA gom quy trình chung vào hệ thống; sếp chốt khác biệt giữa các phòng.",
            ),
            (
                "Order-to-Procurement",
                "Theo dõi nhu cầu giao hàng từ khách và thay đổi sau chốt đơn.",
                "Theo dõi nghĩa vụ thanh toán và điều kiện duyệt chi liên quan đơn mua.",
                "Tạo và theo dõi yêu cầu mua, PO nhà cung cấp, cập nhật trạng thái đặt hàng.",
                "Theo dõi hàng về, chuẩn bị điều kiện nhập kho hoặc bàn giao thẳng.",
                "Huy SA chuẩn hóa trạng thái hệ thống; sếp quyết định ngoại lệ lớn.",
            ),
            (
                "Procurement-to-Delivery",
                "Cập nhật khách hàng về tiến độ giao và xác nhận người nhận / thời điểm nhận.",
                "Chuẩn bị điều kiện xuất hóa đơn, ghi nhận công nợ hoặc doanh thu theo quy định.",
                "Theo dõi hoàn tất mua hàng, xử lý chênh lệch giá/lead time nếu phát sinh.",
                "Kiểm tra tồn, nhập kho, xuất kho, điều phối giao hàng, nghiệm thu và bàn giao chứng từ.",
                "Huy SA theo dõi flow end-to-end; sếp chốt quy trình chuẩn vận hành nếu có xung đột.",
            ),
        ],
        [2.4, 3.2, 3.0, 3.0, 3.0, 4.3],
    )

    add_heading(document, "Phụ lục C. Khung câu hỏi phỏng vấn từng phòng ban")
    add_multi_col_table(
        document,
        ["Câu hỏi", "Mục đích sử dụng"],
        [
            ("Đầu vào hiện nay phòng ban đang nhận là gì?", "Xác định dữ liệu nguồn cần đưa vào hệ thống."),
            ("Đầu ra phòng ban bắt buộc phải tạo là gì?", "Xác định tài liệu/trạng thái cần hệ thống hóa."),
            ("Ai là người phê duyệt hoặc xác nhận bước này?", "Xác định điểm kiểm soát và trách nhiệm."),
            ("Hiện đang dùng file, biểu mẫu hay Excel nào?", "Tìm đúng vật mang dữ liệu cần thay thế hoặc nhập vào hệ thống."),
            ("Lỗi lặp lại nhiều nhất là gì?", "Ưu tiên đúng pain point cần giải quyết bằng hệ thống."),
            ("Điểm nào hiện vẫn phải làm tay hoặc đối chiếu Excel?", "Xác định bước cần số hóa trước trong phase ERP tiếp theo."),
        ],
        [8.2, 10.4],
    )

    add_heading(document, "Phụ lục D. Danh sách module hiện có trong hệ thống")
    add_multi_col_table(
        document,
        ["Nhóm module", "Mức sẵn sàng"],
        [
            ("Dashboard / Reports / Activity", "Đã triển khai"),
            ("Accounts / Contacts / Leads", "Đã triển khai"),
            ("Products / Suppliers / Supplier Quotes", "Đã triển khai"),
            ("Pricing / Quotation / PDF / Status", "Đã triển khai"),
            ("Users / Settings / i18n / Import-Export", "Đã triển khai"),
            ("Sales Orders / ERP outbox", "Đã có nền"),
            ("Projects / Tasks / Ops Overview / Gantt / Staff", "Đã có nền"),
            ("Chat / Notifications", "Đã có nền"),
        ],
        [9.0, 9.6],
    )

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(8)
    run = note.add_run("Kết luận: Dự án nên được triển khai theo mô hình mini app chạy nội bộ, với mục tiêu demo được lát cắt end-to-end Sales-to-Delivery ngay trong tháng 4, sau đó tiếp tục ổn định và mở rộng thành ERP tổng thể theo từng giai đoạn.")
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(10.5)

    return document


if __name__ == "__main__":
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)
