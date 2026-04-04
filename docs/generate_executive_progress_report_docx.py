from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_PATH = r"C:\Users\dinghuy\OneDrive - HUYNH THY GROUP\Antigravity Workspace\htc-erp\docs\bao-cao-executive-crm-docs-code-2026-03-31.docx"

PRIMARY_COLOR = "1F4E78"
LIGHT_FILL = "D9EAF7"


def set_font(run, size=10.5, bold=False, color=None, italic=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, *, bold=False, size=10.0, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table, widths):
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Cm(widths[idx])
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(3)


def add_title(document, title_text, subtitle_text):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title_text)
    set_font(r, size=16, bold=True)

    s = document.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run(subtitle_text)
    set_font(sr, size=10.5, italic=True)


def add_heading(document, text, level=1):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, size=13 if level == 1 else 11.5, bold=True, color=PRIMARY_COLOR)


def add_paragraph(document, text):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r)


def add_bullets(document, items):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_font(r)


def add_two_col_table(document, rows, widths=(5.0, 11.8)):
    table = document.add_table(rows=1, cols=2)
    style_table(table, widths)
    headers = table.rows[0].cells
    set_cell_text(headers[0], "Hạng mục", bold=True, color="FFFFFF")
    set_cell_text(headers[1], "Nội dung", bold=True, color="FFFFFF")
    shade_cell(headers[0], PRIMARY_COLOR)
    shade_cell(headers[1], PRIMARY_COLOR)
    for left, right in rows:
        row = table.add_row().cells
        set_cell_text(row[0], left, bold=True)
        set_cell_text(row[1], right)
    return table


def add_multi_col_table(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    style_table(table, widths)
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color="FFFFFF")
        shade_cell(table.rows[0].cells[idx], PRIMARY_COLOR)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            set_cell_text(cells[idx], value)
    return table


def build_document():
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)

    add_title(
        document,
        "BÁO CÁO EXECUTIVE CRM HUỲNH THY",
        "Tổng hợp từ tài liệu dự án và hiện trạng mã nguồn đến ngày 31/03/2026",
    )

    add_heading(document, "1. Tóm tắt điều hành")
    add_two_col_table(
        document,
        [
            (
                "Bản chất hệ thống",
                "Sản phẩm đang được chuẩn hóa theo hướng CRM gắn ERP, bám revenue flow từ lead đến quotation, approval, project/workspace, task và ERP handoff.",
            ),
            (
                "Hiện trạng chung",
                "Dự án không còn ở trạng thái MVP rời rạc. Hiện đã có trục sản phẩm, kiến trúc, API, role/workspace model, và nhiều phần đã hiện thực trực tiếp trong mã nguồn.",
            ),
            (
                "Cách đọc báo cáo",
                "Báo cáo tách 3 mức evidence: đã định hình trong tài liệu, đã hiện thực trong mã nguồn, và đã có bằng chứng kiểm chứng bằng test/UAT.",
            ),
        ],
    )

    add_heading(document, "2. Những gì đã làm được")
    add_multi_col_table(
        document,
        ["Mức evidence", "Nội dung chính", "Nguồn bằng chứng"],
        [
            (
                "Đã định hình trong tài liệu",
                "Đã xác lập product scope, kiến trúc modular monolith, hướng API /api/v1, role-based workspace, cross-functional model, và phương thức phát triển bằng AI theo spec -> task -> code -> test -> UAT.",
    "docs/crm-overview.md; docs/product/product-spec.md; docs/index.md; docs/ai-delivery-governance-plan.md; docs/architecture/overview.md; docs/api/api-catalog.md; docs/role-permission-matrix-plan.md; docs/cross-functional-v1-plan.md",
            ),
            (
                "Đã có trong mã nguồn",
                "Frontend đã có route thực cho Home, My Work, Inbox, Approvals, Leads, Customers, Products, Suppliers, Quotations, Projects, Tasks, ERP Orders, Reports, EventLog, Users, Settings, Support và Ops. Backend đã có module thực cho auth, crm, quotations, projects, tasks, erp, collaboration, products, suppliers, supplier-quotes, sales-orders, users. SQLite bootstrap đã tách riêng; backend entrypoint đã rút gọn qua src/app.ts.",
                "frontend/src/app.tsx; backend/src/modules/*; backend/src/app.ts; backend/sqlite-db.ts",
            ),
            (
                "Đã có bằng chứng kiểm chứng",
                "Backend đang có test gates cho auth, workspace, pricing, quotation-create-flow, revenue-flow-contracts, api-v1-alias, work-hub, task-view-presets và db-init. Frontend có test:core, test:noncore, test:ux:contracts, typecheck và build. Project Workspace đã có UAT report riêng với evidence rõ nhất ở lớp nghiệp vụ.",
                "backend/package.json; frontend/package.json; docs/project-workspace-uat-report-2026-03-26.md",
            ),
        ],
        [3.2, 8.9, 5.5],
    )

    add_heading(document, "3. Cách thức hoạt động")
    add_bullets(
        document,
        [
            "UI được tổ chức theo role-based home + shared workspace + inbox/approvals, giúp các vai trò làm việc trên cùng một record trung tâm thay vì tách rời theo file Excel hoặc app riêng.",
            "Dữ liệu và workflow bám theo revenue flow, từ đầu vào khách hàng/sản phẩm đến báo giá, approval, handoff và ERP outbox.",
            "Backend đang được tổ chức theo modular monolith để dễ kiểm soát phạm vi thay đổi, giảm chồng chéo và phù hợp với cách giao việc nhỏ cho AI.",
            "ERP integration được định hướng theo outbox + retry + idempotency, tách khỏi flow đồng bộ để tăng khả năng kiểm soát và audit.",
            "Cách triển khai đã được đưa vào khuôn làm việc rõ hơn: spec ngắn, task rõ acceptance criteria, code, test, rồi UAT; giảm dần kiểu làm ad-hoc theo trí nhớ.",
        ],
    )

    add_heading(document, "4. Hướng đi tiếp theo")
    add_multi_col_table(
        document,
        ["Hướng đi", "Mục tiêu", "Cách mô tả trạng thái"],
        [
            (
                "Khóa core revenue flow",
                "Tiếp tục tập trung vào flow chính từ quotation đến handoff, project/workspace, task và ERP handoff trước khi mở rộng các nhánh ngoài lõi.",
                "Định hướng / bước tiếp theo",
            ),
            (
                "Hoàn thiện module ownership và API normalization",
                "Đẩy mạnh true module ownership, API normalization theo /api/v1, workspace/approval flow và role model thống nhất giữa docs, code và UI.",
                "Đang chuẩn hóa",
            ),
            (
                "Nâng độ kiểm chứng và readiness",
                "Tiếp tục tăng độ phủ test/UAT, giữ verification gates bắt buộc và nâng readiness cho rollout lớn hơn khi core flow đủ ổn định.",
                "Bước tiếp theo",
            ),
        ],
        [4.0, 9.6, 4.4],
    )

    appendix = document.add_section(WD_SECTION_START.NEW_PAGE)
    appendix.page_width = Cm(21)
    appendix.page_height = Cm(29.7)
    appendix.top_margin = Cm(1.4)
    appendix.bottom_margin = Cm(1.4)
    appendix.left_margin = Cm(1.5)
    appendix.right_margin = Cm(1.5)

    add_heading(document, "Phụ lục. Nguồn bằng chứng đã dùng")
    add_paragraph(document, "Tài liệu và mã nguồn được ưu tiên để soạn báo cáo này:")
    add_bullets(
        document,
        [
            "docs/crm-overview.md",
            "docs/product/product-spec.md",
    "docs/index.md",
    "docs/ai-delivery-governance-plan.md",
            "docs/architecture/overview.md",
            "docs/api/api-catalog.md",
    "docs/role-permission-matrix-plan.md",
    "docs/cross-functional-v1-plan.md",
            "docs/project-workspace-uat-report-2026-03-26.md",
            "README.md",
            "backend/package.json; frontend/package.json",
            "frontend/src/app.tsx; backend/src/app.ts; backend/sqlite-db.ts",
        ],
    )
    add_paragraph(
        document,
        "Nguyên tắc diễn đạt: nếu một nội dung chỉ mới có trong plan/docs mà chưa có evidence code/test, báo cáo gọi đúng là định hướng, bước tiếp theo, hoặc đang chuẩn hóa; không mô tả như đã hoàn tất.",
    )

    return document


if __name__ == "__main__":
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)
